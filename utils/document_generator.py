"""
公文文档生成器 —— 一键生成符合 GB/T 9704-2012 格式的 .docx 文件。

使用方式：
    from utils.document_generator import write_docx

    write_docx(
        title='[单位名称]关于[事项]的通知',
        subtitle='',
        author='[单位名称]　[姓名]',
        date='XXXX年X月X日',
        greeting='各有关单位：',
        body=[
            '[正文第一段……]',
            '[正文第二段……]',
        ],
        closing='特此通知。',
        category='0001——[类别名称]',
        event_folder='XXXXXXXX——[事项名称]',
    )
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from utils.settings import ROOT

# === 常量 ===
RED = RGBColor(0xFF, 0x00, 0x00)
INDENT = Pt(32)
LQ = '\u201c'
RQ = '\u201d'
LS = '\u2018'
RS = '\u2019'

_PH_RE = re.compile(r'(\[.*?\])')

# === 标点归一化 ===
# 半角 → 全角映射
_ASCII_TO_FULLWIDTH = str.maketrans({
    ',': '\uff0c',
    '.': '\u3002',
    ':': '\uff1a',
    ';': '\uff1b',
    '?': '\uff1f',
    '!': '\uff01',
    '(': '\uff08',
    ')': '\uff09',
})
# 保护模式：保留这些语境下的半角标点
_PUNCT_PROTECT_PATTERNS = [
    r'\d+[.,:]\d+(?:\s*\S*)?',        # 数值: 1.5, 3,000, 10:30
    r'https?://[^\s<>"\']+',           # URL
    r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',  # 邮箱
    r'(?:www\.)[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',             # 域名: www.xxx.xxx
    r'[a-zA-Z0-9-]+\.[a-zA-Z0-9-]+\.[a-zA-Z]{2,}',       # 多级域名: xxx.xxx.xxx
]


def _normalize_punctuation(text: str) -> str:
    """将西文半角标点转为中文全角，同时保护数值/URL/英文语境下的标点。"""
    if not text:
        return text

    # Step 1: 保护数值/URL/邮箱中的标点（用占位符替换以免被转换）
    protected = {}
    counter = [0]

    def _save(m):
        counter[0] += 1
        key = f'\x00P{counter[0]:04d}\x00'
        protected[key] = m.group(0)
        return key

    for pat in _PUNCT_PROTECT_PATTERNS:
        text = re.sub(pat, _save, text)

    # Step 2: ASCII 引号 → 中文引号（双遍扫描，支持嵌套）
    # Pass 1: 双引号 "…" → \u201c…\u201d
    result = []
    toggle = False
    for ch in text:
        if ch == '"':
            result.append('\u201c' if not toggle else '\u201d')
            toggle = not toggle
        else:
            result.append(ch)
    text = ''.join(result)
    if toggle:  # 奇数个双引号：最后一个是左引号
        text = text[:-1] + '\u201c'

    # Pass 2: 单引号 '…' → \u2018…\u2019
    result = []
    toggle = False
    for ch in text:
        if ch == "'":
            result.append('\u2018' if not toggle else '\u2019')
            toggle = not toggle
        else:
            result.append(ch)
    text = ''.join(result)
    if toggle:  # 奇数个单引号：最后一个是左引号
        text = text[:-1] + '\u2018'

    # Step 3: 剩余 ASCII 标点统一转全角
    text = text.translate(_ASCII_TO_FULLWIDTH)

    # Step 4: 恢复被保护的内容
    for key, val in protected.items():
        text = text.replace(key, val)

    return text


def _set_run_font(run, font_name, font_size, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color


def _add_para(doc, text, font_name='仿宋_GB2312', font_size=Pt(16), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.widow_control = False
    p.alignment = alignment
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if text == '':
        return p
    text = _normalize_punctuation(text)
    parts = _PH_RE.split(text)
    for part in parts:
        if part == '':
            continue
        is_ph = _PH_RE.fullmatch(part) is not None
        run = p.add_run('xx' if is_ph else part)
        _set_run_font(run, font_name, font_size, bold, color=RED if is_ph else None)
    return p





def _resolve_target_dir(base_dir: Path, year: str, category: str,
                        event_folder: str | None,
                        fallback_name: str) -> Path:
    target = base_dir / year / category
    if event_folder:
        if target.exists():
            matching = [d for d in target.iterdir()
                        if d.is_dir() and d.name == event_folder]
            if matching:
                return matching[0]
        return target / event_folder
    return target / fallback_name


def _resolve_output_paths(filename_stem: str, category: str,
                          year: str | None = None,
                          event_folder: str | None = None) -> tuple[Path, Path, Path]:

    if year is None:
        year = str(datetime.now().year)

    timestamp = datetime.now().strftime('%Y%m%d')
    fallback_dir = f'{timestamp}——{filename_stem}'

    docx_dir = _resolve_target_dir(
        ROOT / 'output', year, category, event_folder, fallback_dir)
    docx_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%H%M')
    docx_name = f'{datetime.now().strftime("%Y%m%d")}——{filename_stem}（{timestamp}）.docx'
    docx_path = docx_dir / docx_name

    archive_dir = _resolve_target_dir(
        ROOT / 'knowledge-base' / 'archive', year, category,
        event_folder, docx_dir.name)
    archive_dir.mkdir(parents=True, exist_ok=True)

    txt_name = docx_name.replace('.docx', '.txt')
    txt_path = archive_dir / txt_name

    return docx_dir, docx_path, txt_path


def write_docx(
    title: str = '',
    subtitle: str = '',
    author: str = '',
    date: str = '',
    greeting: str = '',
    body: list[str] | None = None,
    closing: str = '',
    doc_number: str = '',
    category: str = '0004——模板',
    year: str | None = None,
    event_folder: str | None = None,
    filename_stem: str = '公文',
) -> Path:
    """一键生成标准格式公文 .docx 并归档 .txt。

    Args:
        title: 公文标题（方正小标宋简体 2号，居中）
        subtitle: 副标题（同上字体）
        author: 作者/讲话人（居中）
        date: 日期（居中）
        greeting: 称呼语，如"同志们："（顶格）
        body: 正文段落列表（首行缩进 2 字符）
        closing: 结束语，如"谢谢大家！"（首行缩进 2 字符）
        doc_number: 发文字号，如"[发文机关代字]发〔2026〕15号"（仿宋_GB2312，居中，位于标题前）
        category: 工作类别，如 "0006——[类别名称]"
        year: 年份，默认当前年份
        event_folder: 已有事项文件夹名，如 "20260401——宣讲比赛活动"
        filename_stem: 文件名主干

    Returns:
        生成的 .docx 文件路径
    """
    if body is None:
        body = []

    # 清理文件名中的非法字符
    safe_stem = re.sub(r'[\\/:*?"<>|]', '_', filename_stem)

    docx_dir, docx_path, txt_path = _resolve_output_paths(
        safe_stem, category, year, event_folder)

    # === 构建 Word 文档 ===
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    pf = style.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = False

    # 发文字号（位于标题前，居中）
    if doc_number:
        _add_para(doc, doc_number, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 标题前空行
    if title:
        _add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 标题
    if title:
        _add_para(doc, title, '方正小标宋简体', Pt(22),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 副标题
    if subtitle:
        _add_para(doc, subtitle, '方正小标宋简体', Pt(22),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 作者
    if author:
        _add_para(doc, author, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 日期
    if date:
        _add_para(doc, date, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 正文前空行
    _add_para(doc, '')

    # 称呼
    if greeting:
        _add_para(doc, greeting, first_line_indent=None)

    # 正文
    for para in body:
        _add_para(doc, para, first_line_indent=INDENT)

    # 结束语
    if closing:
        _add_para(doc, closing, first_line_indent=INDENT)

    # 保存 .docx
    doc.save(str(docx_path))
    print(f'[OK] .docx → {docx_path}')

    # 提取纯文本保存为 .txt 到知识库（与 docx 保持标点一致）
    txt_lines = []
    if doc_number:
        txt_lines.append(_normalize_punctuation(doc_number))
    if title:
        txt_lines.append('')
    for s in (title, subtitle, author, date):
        if s:
            txt_lines.append(_normalize_punctuation(s))
    txt_lines.append('')
    if greeting:
        txt_lines.append(_normalize_punctuation(greeting))
    for para in body:
        txt_lines.append(_normalize_punctuation(para))
    if closing:
        txt_lines.append(_normalize_punctuation(closing))

    txt_path.write_text('\n'.join(txt_lines), encoding='utf-8')
    print(f'[OK] .txt → {txt_path}')

    return docx_path
