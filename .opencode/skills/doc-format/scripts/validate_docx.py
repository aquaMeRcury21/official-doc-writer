"""
公文格式合规校验器。

校验生成的 .docx 是否符合 GB/T 9704-2012 标准。

用法：
  # 校验单个文件
  python scripts/validate_docx.py 路径/文档.docx

  # 校验整个目录下的所有 .docx
  python scripts/validate_docx.py 路径/目录/

  # 静默模式（只输出 pass/fail）
  python scripts/validate_docx.py 路径/文档.docx --quiet
"""

import argparse
import os
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Cm, Pt, Emu


# ============================================================
# 校验规则（GB/T 9704-2012）
# ============================================================

RULES = [
    {
        'id': 'PAGE_SIZE',
        'label': '纸张为 A4（210mm×297mm）',
        'check': lambda doc: _check_page_size(doc),
        'critical': True,
    },
    {
        'id': 'TOP_MARGIN',
        'label': '上边距 3.7cm ±1mm',
        'check': lambda doc: _check_margin(doc, 'top', Cm(3.7)),
        'critical': True,
    },
    {
        'id': 'BOTTOM_MARGIN',
        'label': '下边距 3.5cm ±1mm',
        'check': lambda doc: _check_margin(doc, 'bottom', Cm(3.5)),
        'critical': True,
    },
    {
        'id': 'LEFT_MARGIN',
        'label': '左边距 2.8cm ±1mm',
        'check': lambda doc: _check_margin(doc, 'left', Cm(2.8)),
        'critical': True,
    },
    {
        'id': 'RIGHT_MARGIN',
        'label': '右边距 2.6cm ±1mm',
        'check': lambda doc: _check_margin(doc, 'right', Cm(2.6)),
        'critical': True,
    },
    {
        'id': 'DEFAULT_FONT',
        'label': '默认正文字体 仿宋_GB2312、3号(16pt)',
        'check': lambda doc: _check_default_font(doc),
        'critical': True,
    },
    {
        'id': 'LINE_SPACING',
        'label': '行距固定值 28pt',
        'check': lambda doc: _check_line_spacing(doc),
        'critical': False,
    },
    {
        'id': 'WIDOW_CONTROL',
        'label': '孤行控制已关闭（widow_control=False）',
        'check': lambda doc: _check_widow_control(doc),
        'critical': False,
    },
    {
        'id': 'NO_BOLD_HEADINGS',
        'label': '标题未使用加粗（仅靠字体区分层级）',
        'check': lambda doc: _check_no_bold_headings(doc),
        'critical': False,
    },
    {
        'id': 'FIRST_LINE_INDENT',
        'label': '正文段落首行缩进 2 字符（Pt(32)）',
        'check': lambda doc: _check_first_line_indent(doc),
        'critical': False,
    },
    {
        'id': 'TITLE_FONT',
        'label': '标题使用方正小标宋简体',
        'check': lambda doc: _check_title_font(doc),
        'critical': False,
    },
    {
        'id': 'FULLWIDTH_PUNCTUATION',
        'label': '正文使用全角标点（无西文半角引号）',
        'check': lambda doc: _check_fullwidth_punctuation(doc),
        'critical': False,
    },
]


def _emu_to_cm(emu):
    """EMU -> cm"""
    return emu / 914400 * 2.54 if emu else 0


def _check_page_size(doc):
    for s in doc.sections:
        w = _emu_to_cm(s.page_width)
        h = _emu_to_cm(s.page_height)
        if abs(w - 21.0) > 0.3 or abs(h - 29.7) > 0.3:
            return False, f'实际: {w:.1f}×{h:.1f}cm'
    return True, None


def _check_margin(doc, attr, expected):
    actual = getattr(doc.sections[0], f'{attr}_margin', None)
    if actual is None:
        return False, '未设置'
    actual_cm = _emu_to_cm(actual)
    exp_cm = _emu_to_cm(expected)
    if abs(actual_cm - exp_cm) > 0.15:
        return False, f'实际: {actual_cm:.2f}cm, 期望: {exp_cm:.2f}cm'
    return True, None


def _check_default_font(doc):
    style = doc.styles['Normal']
    issues = []
    ea = style.element.rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
    if ea and '仿宋' not in ea:
        issues.append(f'中文字体: {ea}')
    if style.font.size and abs(style.font.size.pt - 16) > 1:
        issues.append(f'字号: {style.font.size.pt}pt')
    if abs(style.font.size.pt - 16) > 1:
        issues.append(f'西文字体字号: {style.font.size.pt}pt')
    return (not issues), ('; '.join(issues) if issues else None)


def _check_line_spacing(doc):
    pf = doc.styles['Normal'].paragraph_format
    ls = pf.line_spacing
    if ls is None:
        return False, '未设置'
    if abs(ls - Pt(28)) > 2:
        return False, f'实际: {ls}pt'
    return True, None


def _check_widow_control(doc):
    pf = doc.styles['Normal'].paragraph_format
    if pf.widow_control is not False:
        return False, 'widow_control 未关闭'
    return True, None


def _check_no_bold_headings(doc):
    bold_count = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold:
                bold_count += 1
    if bold_count > 3:
        return False, f'发现 {bold_count} 处加粗（>=4 则可能异常）'
    return True, None


def _check_first_line_indent(doc):
    """抽查非标题段落的缩进"""
    no_indent = 0
    total_body = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 跳过明显是标题/版头的段落（居中、大字号、顶格称呼）
        if p.alignment == 1:  # CENTER
            continue
        total_body += 1
        fi = p.paragraph_format.first_line_indent
        if fi is None or fi < Pt(20):
            no_indent += 1
    if no_indent > total_body * 0.5:
        return False, f'超 50% 正文段落无首行缩进（{no_indent}/{total_body}）'
    return True, None


def _check_title_font(doc):
    for p in doc.paragraphs:
        if p.alignment == 1:  # CENTER
            for r in p.runs:
                ea = r._element.rPr.rFonts.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia'
                ) if r._element.rPr is not None and r._element.rPr.rFonts is not None else None
                if ea and '小标宋' not in ea:
                    return False, f'标题行字体: {ea}'
    return True, None


def _check_fullwidth_punctuation(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text and '\u201c' not in r.text:
                if '"' in r.text:
                    return False, f'半角引号残留: "{r.text[:50]}..."'
    return True, None


# ============================================================
# 输出
# ============================================================

def validate(file_path, quiet=False):
    """校验单个 .docx 文件"""
    try:
        doc = Document(file_path)
    except Exception as e:
        return {
            'file': file_path,
            'pass': False,
            'error': f'无法打开文档: {e}',
            'results': [],
            'critical_failures': 1,
        }

    results = []
    critical_failures = 0

    for rule in RULES:
        passed, detail = rule['check'](doc)
        results.append({
            'id': rule['id'],
            'label': rule['label'],
            'passed': passed,
            'detail': detail,
            'critical': rule['critical'],
        })
        if not passed and rule['critical']:
            critical_failures += 1

    overall = critical_failures == 0

    if not quiet:
        status = 'PASS' if overall else 'FAIL'
        print(f'\n[{status}] {os.path.basename(file_path)}')
        print('-' * 60)
        for r in results:
            icon = '✓' if r['passed'] else ('✗' if r['critical'] else '△')
            detail = f' — {r["detail"]}' if r['detail'] else ''
            print(f'  {icon} {r["label"]}{detail}')
        print('-' * 60)
        passed_n = sum(1 for r in results if r['passed'])
        critical_n = sum(1 for r in results if r['critical'])
        critical_passed = sum(1 for r in results if r['passed'] and r['critical'])
        print(f'  通过: {passed_n}/{len(results)} （关键项: {critical_passed}/{critical_n}）')

    return {
        'file': file_path,
        'pass': overall,
        'error': None,
        'results': results,
        'critical_failures': critical_failures,
    }


def main():
    parser = argparse.ArgumentParser(description='公文格式合规校验器')
    parser.add_argument('path', help='.docx 文件或目录')
    parser.add_argument('--quiet', action='store_true', help='静默模式')
    args = parser.parse_args()

    path = args.path
    if not os.path.exists(path):
        print(f'[ERROR] 路径不存在: {path}')
        sys.exit(1)

    files = []
    if os.path.isfile(path) and path.lower().endswith('.docx'):
        files = [path]
    elif os.path.isdir(path):
        for root, _, filenames in os.walk(path):
            for f in filenames:
                if f.lower().endswith('.docx'):
                    files.append(os.path.join(root, f))
    else:
        print(f'[ERROR] 不支持的路径: {path}')
        sys.exit(1)

    if not files:
        print('[WARN] 未找到 .docx 文件')
        sys.exit(0)

    print(f'找到 {len(files)} 个 .docx 文件，开始校验...')

    all_pass = True
    results = []
    for f in files:
        r = validate(f, quiet=args.quiet)
        results.append(r)
        if not r['pass']:
            all_pass = False

    if not args.quiet and len(files) > 1:
        passed = sum(1 for r in results if r['pass'])
        print(f'\n{"=" * 60}')
        print(f'汇总: {passed}/{len(files)} 通过')
        for r in results:
            icon = '✓' if r['pass'] else '✗'
            print(f'  {icon} {os.path.basename(r["file"])}')
        print(f'{"=" * 60}')

    sys.exit(0 if all_pass else 1)


if __name__ == '__main__':
    main()
